"""PDF/DOCX → plain text for resumes (Phase 14).

Only the two formats the spec names are accepted, verified by content
sniffing rather than the client's Content-Type (browsers send
application/octet-stream for .docx). Parsing is CPU-bound library code, so it
runs in a worker thread under a timeout, and any library failure becomes a
422 — a malformed upload is user input, not a server fault.

Phase 18: archives are inspected before decompression (a 5 MB DOCX can
expand to gigabytes) and PDFs are capped by page count, because the timeout
only cancels the await — it cannot stop a thread mid-parse."""

import asyncio
import io
import re
import zipfile
from pathlib import PurePosixPath
from typing import Literal

from app.core.exceptions import InvalidInputError

ResumeKind = Literal["pdf", "docx"]

MIME_TYPES: dict[ResumeKind, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# A résumé is a few pages. These are generous caps, not typical sizes.
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
MAX_DOCX_MEMBERS = 2000
MAX_PDF_PAGES = 100

_UNSUPPORTED = "Only PDF and DOCX resumes are supported."
_NO_TEXT = "No extractable text found — scanned PDFs (OCR) are not supported."
_UNREADABLE = "Could not read this file."
_TOO_COMPLEX = "This file is too large or complex to process as a resume."
_ENCRYPTED = "Password-protected PDFs are not supported — remove the password and re-upload."

# Postgres rejects NUL in text/jsonb; PDF extraction also leaks stray controls.
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_SPACES_RE = re.compile(r"[ \t\f\v]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")


def sanitize_text(value: str) -> str:
    cleaned = _CONTROL_RE.sub("", value).replace("\r\n", "\n").replace("\r", "\n")
    cleaned = _SPACES_RE.sub(" ", cleaned)
    lines = [line.strip() for line in cleaned.split("\n")]
    return _BLANK_LINES_RE.sub("\n\n", "\n".join(lines)).strip()


def sanitize_json(value: object) -> object:
    """Recursively strips control characters from every string in an LLM
    result before it is written to a jsonb column."""
    if isinstance(value, str):
        return _CONTROL_RE.sub("", value)
    if isinstance(value, list):
        return [sanitize_json(item) for item in value]
    if isinstance(value, dict):
        return {str(key): sanitize_json(item) for key, item in value.items()}
    return value


def sniff_kind(filename: str, head: bytes) -> ResumeKind:
    """Extension AND magic bytes must agree; anything else is a 422."""
    suffix = PurePosixPath(filename or "").suffix.casefold()
    if suffix == ".pdf" and b"%PDF-" in head[:1024]:
        return "pdf"
    if suffix == ".docx" and head[:4] == b"PK\x03\x04":
        return "docx"
    raise InvalidInputError(_UNSUPPORTED)


def check_docx_archive(data: bytes) -> None:
    """Rejects decompression bombs using the central directory alone — no
    member is inflated here."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            members = archive.infolist()
            if "word/document.xml" not in {member.filename for member in members}:
                raise InvalidInputError(_UNSUPPORTED)
            if len(members) > MAX_DOCX_MEMBERS:
                raise InvalidInputError(_TOO_COMPLEX)
            total = sum(member.file_size for member in members)
            if total > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise InvalidInputError(_TOO_COMPLEX)
            compressed = max(1, sum(member.compress_size for member in members))
            if total / compressed > MAX_DOCX_COMPRESSION_RATIO:
                raise InvalidInputError(_TOO_COMPLEX)
    except zipfile.BadZipFile as exc:
        raise InvalidInputError(_UNREADABLE) from exc


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    # Owner-password-only PDFs open with an empty user password; anything
    # else is genuinely locked and unreadable for us.
    if reader.is_encrypted and not reader.decrypt(""):
        raise InvalidInputError(_ENCRYPTED)
    if len(reader.pages) > MAX_PDF_PAGES:
        raise InvalidInputError(_TOO_COMPLEX)
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    check_docx_archive(data)

    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cell for cell in cells if cell))
    return "\n".join(parts)


async def extract_text(
    data: bytes, kind: ResumeKind, *, max_chars: int, timeout_seconds: float
) -> str:
    extractor = _extract_pdf if kind == "pdf" else _extract_docx
    try:
        async with asyncio.timeout(timeout_seconds):
            raw = await asyncio.to_thread(extractor, data)
    except InvalidInputError:
        raise
    except TimeoutError as exc:
        raise InvalidInputError(_UNREADABLE) from exc
    except Exception as exc:  # pypdf/python-docx raise many types; all mean "bad file"
        raise InvalidInputError(_UNREADABLE) from exc

    text = sanitize_text(raw)
    if not text:
        raise InvalidInputError(_NO_TEXT)
    return text[:max_chars]
