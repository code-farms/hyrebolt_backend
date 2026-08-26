"""In-memory fakes + file builders for the Phase 14 resume slice."""

import io
import uuid
from dataclasses import dataclass, field
from datetime import UTC, datetime
from types import SimpleNamespace
from typing import Any

from tests.fakes import FakeDB, FakeProfile, FakeProfileRepository

# ── file builders (no fixtures on disk) ─────────────────────────────────────


def make_pdf(lines: list[str]) -> bytes:
    """A minimal but well-formed single-page PDF with a Helvetica text stream,
    so pypdf's extract_text returns the lines."""

    def esc(text: str) -> str:
        return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    content = "BT /F1 12 Tf 72 720 Td 14 TL " + " ".join(f"({esc(line)}) Tj T*" for line in lines) + " ET"
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
            "/Resources << /Font << /F1 5 0 R >> >> >>"
        ),
        f"<< /Length {len(content)} >>\nstream\n{content}\nendstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets: list[int] = []
    for index, body in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{index} 0 obj\n{body}\nendobj\n".encode("latin-1"))
    xref = out.tell()
    out.write(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets:
        out.write(f"{offset:010d} 00000 n \n".encode())
    out.write(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return out.getvalue()


def make_blank_pdf() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def make_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table:
        grid = document.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, value in enumerate(row):
                grid.cell(r, c).text = value
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


RESUME_LINES = [
    "Priya Sharma - Senior Backend Engineer",
    "Experience: Acme Corp, Backend Engineer, 2020 - 2024",
    "Built Python and Postgres services; led migration to Kubernetes.",
    "Education: B.Tech Computer Science, IIT Delhi, 2019",
    "Skills: Python, Django, PostgreSQL, Docker",
]


# ── rows ─────────────────────────────────────────────────────────────────────


def _now() -> datetime:
    return datetime.now(UTC)


@dataclass
class FakeResumeAnalysisRow:
    id: str
    versionId: str
    analysis: dict[str, Any]
    confidence: float | None
    model: str | None
    promptVersion: str | None
    inputTokens: int | None
    outputTokens: int | None
    processedAt: datetime


@dataclass
class FakeVersionRow:
    id: str
    resumeId: str
    versionNumber: int
    fileName: str
    mimeType: str
    fileSize: int
    contentHash: str
    storagePath: str
    extractedText: str
    analysis: FakeResumeAnalysisRow | None = None
    resume: Any = None
    createdAt: datetime = field(default_factory=_now)


@dataclass
class FakeResumeRow:
    id: str
    userId: str
    title: str
    versions: list[FakeVersionRow] = field(default_factory=list)
    createdAt: datetime = field(default_factory=_now)
    updatedAt: datetime = field(default_factory=_now)


@dataclass
class FakeGapRow:
    id: str
    versionId: str
    jobId: str
    analysis: dict[str, Any]
    model: str | None
    promptVersion: str
    processedAt: datetime


# ── repositories ─────────────────────────────────────────────────────────────


class FakeResumeAnalysisRepository:
    def __init__(self) -> None:
        self.rows: dict[str, FakeResumeAnalysisRow] = {}

    async def get_by_version_id(self, version_id: str) -> FakeResumeAnalysisRow | None:
        return self.rows.get(version_id)

    async def upsert_for_version(self, version_id: str, **kwargs: Any) -> FakeResumeAnalysisRow:
        existing = self.rows.get(version_id)
        row = FakeResumeAnalysisRow(
            id=existing.id if existing else uuid.uuid4().hex,
            versionId=version_id,
            analysis=kwargs["analysis"],
            confidence=kwargs["confidence"],
            model=kwargs["model"],
            promptVersion=kwargs["prompt_version"],
            inputTokens=kwargs["input_tokens"],
            outputTokens=kwargs["output_tokens"],
            processedAt=kwargs["processed_at"],
        )
        self.rows[version_id] = row
        return row


class FakeResumeRepository:
    def __init__(
        self, analyses: FakeResumeAnalysisRepository, profiles: FakeProfileRepository | None = None
    ) -> None:
        self.resumes: dict[str, FakeResumeRow] = {}
        self.versions: dict[str, FakeVersionRow] = {}
        self._analyses = analyses
        self._profiles = profiles

    def _hydrate(self, version: FakeVersionRow) -> FakeVersionRow:
        version.analysis = self._analyses.rows.get(version.id)
        version.resume = self.resumes.get(version.resumeId)
        return version

    def _with_versions(self, resume: FakeResumeRow) -> FakeResumeRow:
        resume.versions = sorted(
            (self._hydrate(v) for v in self.versions.values() if v.resumeId == resume.id),
            key=lambda v: v.versionNumber,
            reverse=True,
        )
        return resume

    async def list_for_user(self, user_id: str) -> list[FakeResumeRow]:
        rows = [r for r in self.resumes.values() if r.userId == user_id]
        rows.sort(key=lambda r: r.createdAt, reverse=True)
        return [self._with_versions(r) for r in rows]

    async def get_for_user(self, resume_id: str, user_id: str) -> FakeResumeRow | None:
        row = self.resumes.get(resume_id)
        if row is None or row.userId != user_id:
            return None
        return self._with_versions(row)

    async def create(self, user_id: str, title: str) -> FakeResumeRow:
        row = FakeResumeRow(id=uuid.uuid4().hex, userId=user_id, title=title)
        self.resumes[row.id] = row
        return row

    async def delete(self, resume_id: str) -> None:
        self.resumes.pop(resume_id, None)
        for version_id in [v.id for v in self.versions.values() if v.resumeId == resume_id]:
            self.versions.pop(version_id)
            self._analyses.rows.pop(version_id, None)
        if self._profiles is not None:  # emulate the SetNull relation
            for profile in self._profiles._db.profiles.values():
                if profile.selectedResumeId == resume_id:
                    profile.selectedResumeId = None

    async def create_version(self, *, version_id: str, resume_id: str, **kwargs: Any) -> FakeVersionRow:
        row = FakeVersionRow(
            id=version_id,
            resumeId=resume_id,
            versionNumber=kwargs["version_number"],
            fileName=kwargs["file_name"],
            mimeType=kwargs["mime_type"],
            fileSize=kwargs["file_size"],
            contentHash=kwargs["content_hash"],
            storagePath=kwargs["storage_path"],
            extractedText=kwargs["extracted_text"],
        )
        self.versions[row.id] = row
        return self._hydrate(row)

    async def latest_version(self, resume_id: str) -> FakeVersionRow | None:
        rows = [v for v in self.versions.values() if v.resumeId == resume_id]
        if not rows:
            return None
        return self._hydrate(max(rows, key=lambda v: v.versionNumber))

    async def get_version_for_user(self, version_id: str, user_id: str) -> FakeVersionRow | None:
        version = self.versions.get(version_id)
        if version is None:
            return None
        resume = self.resumes.get(version.resumeId)
        if resume is None or resume.userId != user_id:
            return None
        return self._hydrate(version)

    async def list_storage_paths(self, resume_id: str) -> list[str]:
        return [v.storagePath for v in self.versions.values() if v.resumeId == resume_id]


class FakeResumeGapRepository:
    def __init__(self) -> None:
        self.rows: dict[tuple[str, str], FakeGapRow] = {}

    async def get(self, version_id: str, job_id: str) -> FakeGapRow | None:
        return self.rows.get((version_id, job_id))

    async def upsert(self, version_id: str, job_id: str, **kwargs: Any) -> FakeGapRow:
        row = FakeGapRow(
            id=uuid.uuid4().hex,
            versionId=version_id,
            jobId=job_id,
            analysis=kwargs["analysis"],
            model=kwargs["model"],
            promptVersion=kwargs["prompt_version"],
            processedAt=kwargs["processed_at"],
        )
        self.rows[(version_id, job_id)] = row
        return row


class FakeSkillNames:
    def __init__(self, names: list[str]) -> None:
        self.names = names

    async def list_names(self) -> list[str]:
        return list(self.names)


def make_profiles(user_id: str = "u1", *, skills: list[str] | None = None) -> FakeProfileRepository:
    """A profile repository holding one profile whose skills mirror prisma's
    UserSkill→Skill include shape."""
    db = FakeDB()
    profile = FakeProfile(id=uuid.uuid4().hex, userId=user_id)
    profile.skills = [SimpleNamespace(skill=SimpleNamespace(name=name)) for name in (skills or [])]
    db.profiles[user_id] = profile
    return FakeProfileRepository(db)
